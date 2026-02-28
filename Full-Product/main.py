from __future__ import annotations
"""
Bridge – API Server
B2B SaaS platform for UK financial adviser firms
"""

from fastapi import FastAPI, HTTPException, Query, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from typing import Optional
import os
from dotenv import load_dotenv
load_dotenv()

from mps_data import (
    get_all_mps, get_providers, get_provider, get_mps_by_provider,
    get_mps_by_id, get_platforms, get_investment_styles,
    get_performance_history, filter_mps, get_historical, get_benchmarks,
    get_cost_table,
    get_current_asset_allocation,
)
from insights import (
    get_all_insights, get_insight_by_id, get_insights_by_category,
    get_insight_categories, search_insights,
)
from auth import authenticate, create_session, get_session, destroy_session
from messaging import send_message, get_messages, get_message_by_id
from subscriptions import subscribe, unsubscribe, get_subscriptions, is_subscribed
from preferences import get_preferences, update_preferences, set_subscription_alert

app = FastAPI(
    title="Bridge",
    description="Independent MPS research & oversight platform for UK financial advisers",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


def safe_avg(values):
    """Average that handles None values."""
    clean = [v for v in values if v is not None]
    return round(sum(clean) / len(clean), 2) if clean else 0


def get_current_user(request: Request) -> Optional[dict]:
    """Extract user from session token in cookie or header."""
    token = request.cookies.get("bridge_session") or request.headers.get("X-Session-Token")
    if not token:
        return None
    return get_session(token)


def require_auth(request: Request) -> dict:
    """Dependency: require authenticated user."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(401, "Not authenticated")
    return user


# ─── Auth ──────────────────────────────────────────────────────────────

@app.post("/api/auth/login")
async def login(body: dict):
    email = body.get("email", "").strip()
    password = body.get("password", "")
    if not email or not password:
        raise HTTPException(400, "Email and password required")
    user = authenticate(email, password)
    if not user:
        raise HTTPException(401, "Invalid credentials")
    token = create_session(user)
    from fastapi.responses import JSONResponse
    resp = JSONResponse({"status": "ok", "user": user})
    resp.set_cookie("bridge_session", token, httponly=True, samesite="lax", max_age=86400)
    return resp


@app.post("/api/auth/logout")
async def logout(request: Request):
    token = request.cookies.get("bridge_session") or request.headers.get("X-Session-Token")
    if token:
        destroy_session(token)
    from fastapi.responses import JSONResponse
    resp = JSONResponse({"status": "ok"})
    resp.delete_cookie("bridge_session")
    return resp


@app.get("/api/auth/me")
async def get_me(request: Request):
    user = get_current_user(request)
    if not user:
        return {"authenticated": False}
    subs = get_subscriptions(user["id"])
    return {"authenticated": True, "user": user, "subscriptions": subs}


# ─── Demo Request ─────────────────────────────────────────────────────────

@app.post("/api/demo-request")
async def demo_request(body: dict):
    name = body.get("name", "")
    email = body.get("email", "")
    firm = body.get("firm", "")
    persona = body.get("persona", "Not selected")
    tier = body.get("tier", "Not selected")

    email_subject = f"[Bridge] Demo Request from {name}"
    email_body = (
        f"Name: {name}\n"
        f"Email: {email}\n"
        f"Firm: {firm}\n"
        f"I am: {persona}\n"
        f"Firm size: {tier}"
    )

    if RESEND_API_KEY:
        import httpx
        try:
            async with httpx.AsyncClient() as client:
                r = await client.post(
                    "https://api.resend.com/emails",
                    headers={"Authorization": f"Bearer {RESEND_API_KEY}"},
                    json={
                        "from": "Bridge Demo <onboarding@resend.dev>",
                        "to": [FEEDBACK_EMAIL],
                        "subject": email_subject,
                        "text": email_body,
                    },
                )
                if r.status_code not in (200, 201):
                    print(f"Resend error: {r.text}")
        except Exception as e:
            print(f"Resend error: {e}")
    else:
        print(f"\n--- DEMO REQUEST ---")
        print(email_body)
        print(f"(Set RESEND_API_KEY to send via email)")
        print(f"--------------------\n")

    return {"status": "ok"}

# ─── Messaging ─────────────────────────────────────────────────────────

@app.post("/api/messages")
async def create_message(body: dict, user: dict = Depends(require_auth)):
    subject = body.get("subject", "").strip()
    message_body = body.get("message", "").strip()
    provider_id = body.get("provider_id")
    provider_name = body.get("provider_name")
    if not subject or not message_body:
        raise HTTPException(400, "Subject and message are required")
    msg = send_message(
        user_id=user["id"],
        user_name=user["name"],
        user_firm=user["firm"],
        subject=subject,
        body=message_body,
        provider_id=provider_id,
        provider_name=provider_name,
    )

    # Send email notification to Bridge team
    email_subject = f"[Bridge Message] {subject}"
    if provider_name:
        email_subject += f" — re: {provider_name}"
    email_body = (
        f"From: {user['name']} ({user['email']})\n"
        f"Firm: {user['firm']}\n"
        f"Provider: {provider_name or 'General'}\n\n"
        f"{message_body}"
    )

    if RESEND_API_KEY:
        import httpx
        try:
            async with httpx.AsyncClient() as client:
                r = await client.post(
                    "https://api.resend.com/emails",
                    headers={"Authorization": f"Bearer {RESEND_API_KEY}"},
                    json={
                        "from": "Bridge Messages <onboarding@resend.dev>",
                        "to": [FEEDBACK_EMAIL],
                        "subject": email_subject,
                        "text": email_body,
                    },
                )
                if r.status_code not in (200, 201):
                    print(f"Resend error: {r.text}")
        except Exception as e:
            print(f"Resend error: {e}")
    else:
        print(f"\n--- MESSAGE ---")
        print(f"Subject: {email_subject}")
        print(f"{email_body}")
        print(f"(Set RESEND_API_KEY to send via email)")
        print(f"---------------\n")

    return {"status": "ok", "message": msg}


@app.get("/api/messages")
async def list_messages(user: dict = Depends(require_auth)):
    msgs = get_messages(user["id"])
    return {"count": len(msgs), "messages": msgs}


@app.get("/api/messages/{message_id}")
async def get_message_detail(message_id: str, user: dict = Depends(require_auth)):
    msg = get_message_by_id(message_id, user["id"])
    if not msg:
        raise HTTPException(404, "Message not found")
    return {"message": msg}


# ─── Subscriptions ─────────────────────────────────────────────────────

@app.post("/api/subscriptions")
async def create_subscription(body: dict, user: dict = Depends(require_auth)):
    provider_id = body.get("provider_id", "").strip()
    provider_name = body.get("provider_name", "").strip()
    if not provider_id:
        raise HTTPException(400, "Provider ID required")
    sub = subscribe(user["id"], provider_id, provider_name, user_email=user.get("email", ""))
    return {"status": "ok", "subscription": sub}


@app.delete("/api/subscriptions/{provider_id}")
async def remove_subscription(provider_id: str, user: dict = Depends(require_auth)):
    success = unsubscribe(user["id"], provider_id)
    if not success:
        raise HTTPException(404, "Subscription not found")
    return {"status": "ok"}


@app.get("/api/subscriptions")
async def list_subscriptions(user: dict = Depends(require_auth)):
    subs = get_subscriptions(user["id"])
    return {"count": len(subs), "subscriptions": subs}


@app.get("/api/subscriptions/check/{provider_id}")
async def check_subscription(provider_id: str, user: dict = Depends(require_auth)):
    return {"subscribed": is_subscribed(user["id"], provider_id)}


# ─── Preferences ───────────────────────────────────────────────────────

@app.get("/api/preferences")
async def get_user_preferences(user: dict = Depends(require_auth)):
    prefs = get_preferences(user["id"])
    return {"preferences": prefs}


@app.put("/api/preferences")
async def update_user_preferences(body: dict, user: dict = Depends(require_auth)):
    prefs = update_preferences(user["id"], body)
    return {"status": "ok", "preferences": prefs}


@app.put("/api/preferences/subscription-alert")
async def toggle_subscription_alert(body: dict, user: dict = Depends(require_auth)):
    provider_id = body.get("provider_id", "").strip()
    enabled = body.get("enabled", True)
    if not provider_id:
        raise HTTPException(400, "Provider ID required")
    prefs = set_subscription_alert(user["id"], provider_id, enabled)
    return {"status": "ok", "preferences": prefs}


# ─── Selection Module ──────────────────────────────────────────────────

@app.get("/api/selection/filters")
async def get_filter_options():
    all_mps = get_all_mps()
    risk_ratings = sorted(set(m["risk_rating"] for m in all_mps))
    providers = sorted(set(m["provider"] for m in all_mps))
    # Collect unique fund selection styles from provider data
    all_providers = get_providers()
    styles = sorted(set(p.get("investment_style", "") for p in all_providers.values() if p.get("investment_style")))
    # Collect unique risk rating providers
    rrp = set()
    for prov in all_providers.values():
        for r in prov.get("risk_rating_providers", []):
            rrp.add(r)
    return {
        "platforms": get_platforms(),
        "providers": providers,
        "investment_styles": styles,
        "risk_ratings": risk_ratings,
        "risk_range": {"min": min(risk_ratings), "max": max(risk_ratings)} if risk_ratings else {},
        "risk_rating_providers": sorted(rrp),
    }


@app.get("/api/selection/mps")
async def search_mps(
    risk_min: int = Query(1, ge=1, le=10),
    risk_max: int = Query(10, ge=1, le=10),
    platforms: Optional[str] = None,
    providers: Optional[str] = None,
    ethical_only: bool = False,
    decumulation: bool = False,
    time_horizon: Optional[str] = None,
    max_ocf: Optional[float] = None,
):
    platform_list = platforms.split(",") if platforms else None
    provider_list = providers.split(",") if providers else None

    results = filter_mps(
        risk_min=risk_min, risk_max=risk_max,
        platforms=platform_list, providers=provider_list,
        ethical_only=ethical_only, decumulation=decumulation,
        time_horizon=time_horizon, ocf_max=max_ocf,
    )
    return {"count": len(results), "mps": results}


# ─── Analysis Module ───────────────────────────────────────────────────

@app.get("/api/providers")
async def list_providers():
    providers = get_providers()
    result = []
    for name, data in providers.items():
        portfolios = get_mps_by_provider(name)
        result.append({
            **data,
            "portfolio_count": len(portfolios),
            "risk_range": f"{min(p['risk_rating'] for p in portfolios)}-{max(p['risk_rating'] for p in portfolios)}" if portfolios else "N/A",
            "ocf_range": f"{min(p['ocf'] for p in portfolios):.2f}-{max(p['ocf'] for p in portfolios):.2f}" if portfolios else "N/A",
        })
    return {"providers": result}


@app.get("/api/providers/{provider_id}")
async def get_provider_detail(provider_id: str):
    provider = None
    for name, data in get_providers().items():
        if data["id"] == provider_id:
            provider = data
            break

    if not provider:
        raise HTTPException(404, "Provider not found")

    portfolios = get_mps_by_provider(provider["name"])

    return {
        "provider": provider,
        "portfolios": portfolios,
        "risk_return_data": provider.get("risk_return_data", []),
        "analytics": {
            "avg_ocf": safe_avg([p["ocf"] for p in portfolios]),
            "avg_return_1yr": safe_avg([p.get("return_1yr") for p in portfolios]),
            "avg_return_3yr": safe_avg([p.get("return_3yr") for p in portfolios]),
            "avg_return_5yr": safe_avg([p.get("return_5yr") for p in portfolios]),
            "platform_count": len(set(p for port in portfolios for p in port.get("platforms", []))),
        },
    }
@app.get("/api/mps/{mps_id}")
async def get_mps_detail(mps_id: str):
    mps = get_mps_by_id(mps_id)
    if not mps:
        raise HTTPException(404, "MPS not found")

    provider = get_provider(mps["provider"])
    history = get_performance_history(mps_id, months=36)

    peers = [m for m in get_all_mps()
             if m["risk_rating"] == mps["risk_rating"] and m["id"] != mps_id]

    return {
        "mps": mps,
        "provider": provider,
        "performance_history": history,
        "peer_comparison": {
            "count": len(peers),
            "peers": peers,
            "avg_ocf": safe_avg([p["ocf"] for p in peers]),
            "avg_return_1yr": safe_avg([p.get("return_1yr") for p in peers]),
            "avg_return_3yr": safe_avg([p.get("return_3yr") for p in peers]),
        },
    }


@app.get("/api/mps/{mps_id}/performance")
async def get_mps_performance(mps_id: str, months: int = Query(36, ge=6, le=60)):
    mps = get_mps_by_id(mps_id)
    if not mps:
        raise HTTPException(404, "MPS not found")
    return {"mps_id": mps_id, "history": get_performance_history(mps_id, months)}


@app.get("/api/compare")
async def compare_mps(ids: str = Query(..., description="Comma-separated MPS IDs")):
    id_list = [i.strip() for i in ids.split(",")]
    results = [get_mps_by_id(i) for i in id_list if get_mps_by_id(i)]

    if not results:
        raise HTTPException(404, "No valid MPS found")

    return {
        "count": len(results),
        "mps": results,
    }


# ─── Historical & Benchmarks ──────────────────────────────────────────

@app.get("/api/historical/{key}")
async def get_historical_data(key: str):
    data = get_historical(key)
    if not data:
        raise HTTPException(404, "Historical data not found")
    return data


@app.get("/api/benchmarks")
async def get_benchmark_data():
    return get_benchmarks()


@app.get("/api/costs")
async def get_costs():
    return {"costs": get_cost_table()}


@app.get("/api/current-asset-allocation")
async def get_current_aa():
    return {"asset_allocation": get_current_asset_allocation()}


# ─── Insights Module ──────────────────────────────────────────────────

@app.get("/api/insights")
async def list_insights(
    category: Optional[str] = None,
    search: Optional[str] = None,
):
    if search:
        results = search_insights(search)
    elif category:
        results = get_insights_by_category(category)
    else:
        results = get_all_insights()

    return {
        "count": len(results),
        "insights": results,
        "categories": get_insight_categories(),
    }


@app.get("/api/insights/{insight_id}")
async def get_insight_detail(insight_id: str):
    insight = get_insight_by_id(insight_id)
    if not insight:
        raise HTTPException(404, "Insight not found")
    return {"insight": insight}


# ─── Dashboard ─────────────────────────────────────────────────────────

@app.get("/api/dashboard")
async def get_dashboard():
    all_mps = get_all_mps()
    providers = get_providers()
    insights = get_all_insights()

    return {
        "stats": {
            "total_mps": len(all_mps),
            "total_providers": len(providers),
            "total_platforms": len(get_platforms()),
            "latest_insights": len(insights),
            "market_aum_bn": sum(p.get("aum_bn", 0) for p in providers.values()),
        },
        "recent_insights": insights[:3],
        "provider_overview": [
            {
                "name": name,
                "aum_bn": data.get("aum_bn", 0),
                "style": data.get("investment_style", ""),
                "portfolio_count": len(get_mps_by_provider(name)),
            }
            for name, data in providers.items()
        ],
        "risk_distribution": {
            str(r): len([m for m in all_mps if m["risk_rating"] == r])
            for r in sorted(set(m["risk_rating"] for m in all_mps))
        },
    }


@app.get("/api/health")
async def health():
    return {"status": "healthy", "version": "1.0.0", "platform": "Bridge"}


# ─── Feedback ──────────────────────────────────────────────────────────

FEEDBACK_EMAIL = os.environ.get("FEEDBACK_EMAIL", "feedback@bridge.example.com")
RESEND_API_KEY = os.environ.get("RESEND_API_KEY", "")

@app.post("/api/feedback")
async def submit_feedback(body: dict):
    subject = body.get("subject", "").strip()
    message = body.get("message", "").strip()
    if not subject or not message:
        raise HTTPException(400, "Subject and message are required")

    if RESEND_API_KEY:
        import httpx
        try:
            async with httpx.AsyncClient() as client:
                r = await client.post(
                    "https://api.resend.com/emails",
                    headers={"Authorization": f"Bearer {RESEND_API_KEY}"},
                    json={
                        "from": "Bridge Feedback <onboarding@resend.dev>",
                        "to": [FEEDBACK_EMAIL],
                        "subject": f"[Bridge Feedback] {subject}",
                        "text": message,
                    },
                )
                if r.status_code not in (200, 201):
                    print(f"Resend error: {r.text}")
                    raise HTTPException(500, "Failed to send feedback.")
        except HTTPException:
            raise
        except Exception as e:
            print(f"Resend error: {e}")
            raise HTTPException(500, "Failed to send feedback.")
    else:
        print(f"\n--- FEEDBACK ---")
        print(f"Subject: {subject}")
        print(f"Message: {message}")
        print(f"(Set RESEND_API_KEY env var to send via email)")
        print(f"----------------\n")

    return {"status": "ok", "message": "Feedback received"}

# ─── Consumer Duty Export ────────────────────────────────────────────

@app.post("/api/export/consumer-duty")
async def export_consumer_duty(data: dict):
    """Generate a formatted .docx Consumer Duty Oversight Report."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import io

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x1F, 0x2E)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Title
    title = doc.add_heading("Consumer Duty Oversight Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x2E)

    doc.add_paragraph("")

    # Report Details table
    details = data.get("details", {})
    if details:
        detail_heading = doc.add_heading("Report Details", level=2)
        for run in detail_heading.runs:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.size = Pt(14)

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for key, val in details.items():
            row = table.add_row()
            cell_key = row.cells[0]
            cell_val = row.cells[1]
            cell_key.text = key
            cell_val.text = val
            for paragraph in cell_key.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10.5)
            for paragraph in cell_val.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)

        doc.add_paragraph("")

    # Content sections
    sections = data.get("sections", [])
    for sec in sections:
        sec_title = sec.get("title", "")
        sec_content = sec.get("content", "")

        heading = doc.add_heading(sec_title, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.size = Pt(14)

        # Split content into paragraphs
        for para_text in sec_content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(4)

        doc.add_paragraph("")

    # Footer
    doc.add_paragraph("")
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_line.add_run("Generated by Bridge – Independent MPS Research & Oversight")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.italic = True

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Consumer_Duty_Oversight_Report.docx"},
    )


# ─── Serve Frontend ───────────────────────────────────────────────────

from fastapi.responses import FileResponse

@app.get("/favicon.ico")
async def favicon():
    path = os.path.join(os.path.dirname(__file__), "favicon.ico")
    if os.path.exists(path):
        return FileResponse(path, media_type="image/x-icon")
    raise HTTPException(404)

@app.get("/favicon-32x32.png")
async def favicon_png():
    path = os.path.join(os.path.dirname(__file__), "favicon-32x32.png")
    if os.path.exists(path):
        return FileResponse(path, media_type="image/png")
    raise HTTPException(404)

@app.get("/apple-touch-icon.png")
async def apple_touch_icon():
    path = os.path.join(os.path.dirname(__file__), "apple-touch-icon.png")
    if os.path.exists(path):
        return FileResponse(path, media_type="image/png")
    raise HTTPException(404)

@app.get("/bridge-logo-nav.png")
async def logo_nav():
    path = os.path.join(os.path.dirname(__file__), "bridge-logo-nav.png")
    if os.path.exists(path):
        return FileResponse(path, media_type="image/png")
    raise HTTPException(404)

@app.get("/bridge-logo-dark.png")
async def logo_dark():
    path = os.path.join(os.path.dirname(__file__), "bridge-logo-dark.png")
    if os.path.exists(path):
        return FileResponse(path, media_type="image/png")
    raise HTTPException(404)

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    html_path = os.path.join(os.path.dirname(__file__), "index.html")
    if os.path.exists(html_path):
        with open(html_path, "r") as f:
            return f.read()
    return "<h1>Bridge</h1><p>Frontend not found. See <a href='/docs'>/docs</a></p>"


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
